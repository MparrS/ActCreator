import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from datetime import date
from docx.shared import Pt
from docx2pdf import convert
import openpyxl
import os

def copiar_estilo(origen, destino):
    destino.bold = origen.bold
    destino.italic = origen.italic
    destino.underline = origen.underline
    destino.font.name = origen.font.name
    destino.font.size = origen.font.size

    #Estas son las variables que se van a a remplazar
def modificar_documento(docx_path, nombre, cedula, cpu, monitor, diadema, pin, nombre_archivo):

    # Abrir el documento Word
    doc = Document(docx_path)

    # Obtener la fecha actual
    fecha_actual = date.today().strftime("%d/%m/%Y")

    # Reemplazar texto en párrafos
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if '#fecha' in run.text:
                run.text = run.text.replace('#fecha', fecha_actual)
                copiar_estilo(run, run)
            if '#nombre' in run.text:
                run.text = run.text.replace('#nombre', nombre)
                copiar_estilo(run, run)
            if '#cedula' in run.text:
                run.text = run.text.replace('#cedula', cedula)
                copiar_estilo(run, run)
            if '#cpu' in run.text:
                run.text = run.text.replace('#cpu', cpu)
                copiar_estilo(run, run)
            if '#monitor' in run.text:
                run.text = run.text.replace('#monitor', monitor)
                copiar_estilo(run, run)
            if '#diadema' in run.text:
                run.text = run.text.replace('#diadema', diadema)
                copiar_estilo(run, run)
            if '#pin' in run.text:
                run.text = run.text.replace('#pin', pin)
                copiar_estilo(run, run)

    # Reemplazar texto en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if '#fecha' in run.text:
                            run.text = run.text.replace('#fecha', fecha_actual)
                            copiar_estilo(run, run)
                        if '#nombre' in run.text:
                            run.text = run.text.replace('#nombre', nombre)
                            copiar_estilo(run, run)
                        if '#cedula' in run.text:
                            run.text = run.text.replace('#cedula', cedula)
                            copiar_estilo(run, run)
                        if '#cpu' in run.text:
                            run.text = run.text.replace('#cpu', cpu)
                            copiar_estilo(run, run)
                        if '#monitor' in run.text:
                            run.text = run.text.replace('#monitor', monitor)
                            copiar_estilo(run, run)
                        if '#diadema' in run.text:
                            run.text = run.text.replace('#diadema', diadema)
                            copiar_estilo(run, run)
                        if '#pin' in run.text:
                            run.text = run.text.replace('#pin', pin)
                            copiar_estilo(run, run)

    # Nombre de carpeta en la cual van a ir los words
    carpeta_word = 'Documentos Word'
    if not os.path.exists(carpeta_word):
        os.makedirs(carpeta_word)
    doc.save(os.path.join(carpeta_word, nombre_archivo + '.docx'))

    # Nombre de carpeta en la cual van a ir los PDF
    carpeta_pdf = 'Documentos PDF'
    if not os.path.exists(carpeta_pdf):
        os.makedirs(carpeta_pdf)
    convert(os.path.join(carpeta_word, nombre_archivo + '.docx'), os.path.join(carpeta_pdf, nombre_archivo + '.pdf'))


    # Funcion para la integracion de interfaz grafica
def abrir_ventana():
    ventana = tk.Tk()
    ventana.title("ActCreator")
    ventana.geometry("500x300")

    def seleccionar_archivo_word():
        archivo_word = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        ruta_word.set(archivo_word)

    def seleccionar_archivo_excel():
        archivo_excel = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        ruta_excel.set(archivo_excel)

    def generar_documentos():
        archivo_word = ruta_word.get()
        archivo_excel = ruta_excel.get()

        if not archivo_word or not archivo_excel:
            messagebox.showerror("Error", "Por favor selecciona los archivos Word y Excel.")
            return

        wb = openpyxl.load_workbook(archivo_excel)
        hoja = wb.active

        documentos_generados = []

        total_documentos = hoja.max_row - 1

        for i, fila in enumerate(hoja.iter_rows(min_row=2, values_only=True), start=1):
            nombre, cedula, cpu, monitor, diadema, pin = fila
            nombre_archivo = 'Acta ' + nombre
            modificar_documento(archivo_word, str(nombre), str(cedula), str(cpu), str(monitor), str(diadema), str(pin), nombre_archivo)

            documentos_generados.append(nombre_archivo)

        wb.close()

        # Mensaje después de haber terminado de generar todos los documentos
        messagebox.showinfo("Éxito", f"Se han generado los siguientes documentos: {', '.join(documentos_generados)}")
        
        # Version alterna | messagebox.showinfo("Éxito", f"Se han generado todos los documentos con exito")

    # Variables para almacenar las rutas de los archivos
    ruta_word = tk.StringVar()
    ruta_excel = tk.StringVar()

    # Botones para seleccionar archivos
    select_button_word = tk.Button(ventana, text="Seleccionar archivo Word", command=seleccionar_archivo_word)
    select_button_word.pack(pady=10)

    select_button_excel = tk.Button(ventana, text="Seleccionar archivo Excel", command=seleccionar_archivo_excel)
    select_button_excel.pack(pady=10)

    # Botón para generar los documentos
    generate_button = tk.Button(ventana, text="Generar documentos", command=generar_documentos)
    generate_button.pack(pady=10)

    ventana.mainloop()

abrir_ventana() 