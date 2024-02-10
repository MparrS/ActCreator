from tkinter import Tk, Label, Entry, Button, StringVar
from docx import Document
from datetime import date
from docx.shared import Pt
from docx2pdf import convert

def copiar_estilo(origen, destino):
    destino.bold = origen.bold
    destino.italic = origen.italic
    destino.underline = origen.underline
    destino.font.name = origen.font.name
    destino.font.size = origen.font.size

def modificar_documento(nombre, cedula, cpu, monitor, diadema, pin):
    
    # Ruta del archivo Word original
    archivo_word = r"C:\Users\mserr\Downloads\acta.docx"
    
    # Nombre del archivo modificado
    nombre_archivo = 'Acta' + nombre

    # Abrir el documento Word
    doc = Document(archivo_word)

    # Obtener la fecha actual
    fecha_actual = date.today().strftime("%d/%m/%Y")

    # Reemplazar texto en párrafos
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if '#fecha' in run.text:
                run.text = run.text.replace('#fecha', fecha_actual)
                copiar_estilo(run, run)
            if '#nombre' in run.text:
                run.text = run.text.replace('#nombre', nombre)
                copiar_estilo(run, run)
            if '#cedula' in run.text:
                run.text = run.text.replace('#cedula', cedula)
                copiar_estilo(run, run)
            if '#cpu' in run.text:
                run.text = run.text.replace('#cpu', cpu)
                copiar_estilo(run, run)
            if '#monitor' in run.text:
                run.text = run.text.replace('#monitor', monitor)
                copiar_estilo(run, run)
            if '#diadema' in run.text:
                run.text = run.text.replace('#diadema', diadema)
                copiar_estilo(run, run)
            if '#pin' in run.text:
                run.text = run.text.replace('#pin', pin)
                copiar_estilo(run, run)

    # Reemplazar texto en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if '#fecha' in run.text:
                            run.text = run.text.replace('#fecha', fecha_actual)
                            copiar_estilo(run, run)
                        if '#nombre' in run.text:
                            run.text = run.text.replace('#nombre', nombre)
                            copiar_estilo(run, run)
                        if '#cedula' in run.text:
                            run.text = run.text.replace('#cedula', cedula)
                            copiar_estilo(run, run)
                        if '#cpu' in run.text:
                            run.text = run.text.replace('#cpu', cpu)
                            copiar_estilo(run, run)
                        if '#monitor' in run.text:
                            run.text = run.text.replace('#monitor', monitor)
                            copiar_estilo(run, run)
                        if '#diadema' in run.text:
                            run.text = run.text.replace('#diadema', diadema)
                            copiar_estilo(run, run)
                        if '#pin' in run.text:
                            run.text = run.text.replace('#pin', pin)
                            copiar_estilo(run, run)

    # Guardar el documento modificado
    doc.save(nombre_archivo + '.docx')

    # Convertir el documento modificado a PDF
    convert(nombre_archivo + '.docx', nombre_archivo + '.pdf')

    print("El documento se ha modificado y convertido a PDF correctamente.")

def guardar_datos(nombre, cedula, cpu, monitor, diadema, pin):
    modificar_documento(nombre, cedula, cpu, monitor, diadema, pin)

def abrir_ventana():
    ventana = Tk()
    ventana.title("Ingresar Datos")
    ventana.geometry("300x250")

    nombre_var = StringVar()
    cedula_var = StringVar()
    cpu_var = StringVar()
    monitor_var = StringVar()
    diadema_var = StringVar()
    pin_var = StringVar()

    Label(ventana, text="Nombre:").grid(row=0, column=0)
    Entry(ventana, textvariable=nombre_var).grid(row=0, column=1)
    Label(ventana, text="Cédula:").grid(row=1, column=0)
    Entry(ventana, textvariable=cedula_var).grid(row=1, column=1)
    Label(ventana, text="CPU:").grid(row=2, column=0)
    Entry(ventana, textvariable=cpu_var).grid(row=2, column=1)
    Label(ventana, text="Monitor:").grid(row=3, column=0)
    Entry(ventana, textvariable=monitor_var).grid(row=3, column=1)
    Label(ventana, text="Diadema:").grid(row=4, column=0)
    Entry(ventana, textvariable=diadema_var).grid(row=4, column=1)
    Label(ventana, text="PIN:").grid(row=5, column=0)
    Entry(ventana, textvariable=pin_var).grid(row=5, column=1)

    Button(ventana, text="Aceptar", command=lambda: guardar_datos(nombre_var.get(), cedula_var.get(), cpu_var.get(), monitor_var.get(), diadema_var.get(), pin_var.get())).grid(row=6, column=0, columnspan=2)

    ventana.mainloop()

abrir_ventana()
