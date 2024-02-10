from docx import Document
from datetime import date
from docx.shared import Pt
from docx2pdf import convert
import openpyxl

def copiar_estilo(origen, destino):
    destino.bold = origen.bold
    destino.italic = origen.italic
    destino.underline = origen.underline
    destino.font.name = origen.font.name
    destino.font.size = origen.font.size

def modificar_documento(docx_path, nombre, cedula, cpu, monitor, diadema, pin, nombre_archivo):
    
    # Abrir el documento Word
    doc = Document(docx_path)

    # Obtener la fecha actual
    fecha_actual = date.today().strftime("%d/%m/%Y")

    # Reemplazar texto en párrafos
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if '#fecha' in run.text:
                run.text = run.text.replace('#fecha', fecha_actual)
                copiar_estilo(run, run)
            if '#nombre' in run.text:
                run.text = run.text.replace('#nombre', nombre)
                copiar_estilo(run, run)
            if '#cedula' in run.text:
                run.text = run.text.replace('#cedula', cedula)
                copiar_estilo(run, run)
            if '#cpu' in run.text:
                run.text = run.text.replace('#cpu', cpu)
                copiar_estilo(run, run)
            if '#monitor' in run.text:
                run.text = run.text.replace('#monitor', monitor)
                copiar_estilo(run, run)
            if '#diadema' in run.text:
                run.text = run.text.replace('#diadema', diadema)
                copiar_estilo(run, run)
            if '#pin' in run.text:
                run.text = run.text.replace('#pin', pin)
                copiar_estilo(run, run)

    # Reemplazar texto en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if '#fecha' in run.text:
                            run.text = run.text.replace('#fecha', fecha_actual)
                            copiar_estilo(run, run)
                        if '#nombre' in run.text:
                            run.text = run.text.replace('#nombre', nombre)
                            copiar_estilo(run, run)
                        if '#cedula' in run.text:
                            run.text = run.text.replace('#cedula', cedula)
                            copiar_estilo(run, run)
                        if '#cpu' in run.text:
                            run.text = run.text.replace('#cpu', cpu)
                            copiar_estilo(run, run)
                        if '#monitor' in run.text:
                            run.text = run.text.replace('#monitor', monitor)
                            copiar_estilo(run, run)
                        if '#diadema' in run.text:
                            run.text = run.text.replace('#diadema', diadema)
                            copiar_estilo(run, run)
                        if '#pin' in run.text:
                            run.text = run.text.replace('#pin', pin)
                            copiar_estilo(run, run)

    # Guardar el documento modificado
    doc.save(nombre_archivo + '.docx')

    # Convertir el documento modificado a PDF
    convert(nombre_archivo + '.docx', nombre_archivo + '.pdf')

# Ruta del archivo Word original
archivo_word = r"C:\Users\mserr\Downloads\acta.docx"

# Ruta del archivo Excel con los datos
archivo_excel = r"C:\Users\mserr\Downloads\inventario.xlsx"

# Cargar el archivo Excel
wb = openpyxl.load_workbook(archivo_excel)
hoja = wb.active

# Iterar sobre las filas del archivo Excel
for fila in hoja.iter_rows(min_row=2, values_only=True):
    nombre, cedula, cpu, monitor, diadema, pin = fila
    
    # Nombre del archivo modificado
    nombre_archivo = 'Acta ' + nombre

    # Modificar el documento
    modificar_documento(archivo_word, str(nombre), str(cedula), str(cpu), str(monitor), str(diadema), str(pin), nombre_archivo)



    print(f"El documento {nombre_archivo} se ha modificado y convertido a PDF correctamente.")

# Cerrar el archivo Excel
wb.close()
